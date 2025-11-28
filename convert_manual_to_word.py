#!/usr/bin/env python3
"""
Convert USER_MANUAL.md to Word Document (.docx)
Library Management System - Class XII Project
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document():
    """Convert USER_MANUAL.md to a formatted Word document"""

    # Create a new Document
    doc = Document()

    # Set document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Read the markdown file
    with open('USER_MANUAL.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Split into lines
    lines = content.split('\n')

    in_code_block = False
    in_table = False
    table_obj = None
    table_headers = []

    for i, line in enumerate(lines):
        # Skip empty lines at the start
        if not line.strip() and i < 5:
            continue

        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                # Start of code block - add heading if present
                code_lang = line.strip()[3:].strip()
                if code_lang:
                    p = doc.add_paragraph(f'[Code: {code_lang}]', style='Intense Quote')
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            continue

        if in_code_block:
            # Inside code block
            p = doc.add_paragraph(line, style='Intense Quote')
            if p.runs:
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(9)
            continue

        # Headers
        if line.startswith('# ') and not line.startswith('##'):
            # H1 - Main title
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        elif line.startswith('## '):
            # H2
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('### '):
            # H3
            doc.add_heading(line[4:], level=3)
            continue
        elif line.startswith('#### '):
            # H4
            doc.add_heading(line[5:], level=4)
            continue
        elif line.startswith('##### '):
            # H5
            doc.add_heading(line[6:], level=4)  # Use level 4 for H5
            continue

        # Horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('_' * 80)
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                # Start of table - get headers
                in_table = True
                table_headers = [cell.strip() for cell in line.split('|')[1:-1]]
                table_obj = doc.add_table(rows=1, cols=len(table_headers))
                table_obj.style = 'Light Grid Accent 1'

                # Add headers
                header_cells = table_obj.rows[0].cells
                for idx, header in enumerate(table_headers):
                    header_cells[idx].text = header
                    # Make header bold
                    for paragraph in header_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                continue
            elif line.strip().startswith('|--') or line.strip().startswith('| --'):
                # Table separator line - skip
                continue
            else:
                # Table data row
                cells_data = [cell.strip() for cell in line.split('|')[1:-1]]
                if len(cells_data) == len(table_headers):
                    row_cells = table_obj.add_row().cells
                    for idx, cell_data in enumerate(cells_data):
                        row_cells[idx].text = cell_data
                continue
        else:
            if in_table:
                in_table = False
                table_obj = None
                doc.add_paragraph()  # Add space after table

        # Screenshot placeholders
        if line.strip().startswith('**[SCREENSHOT'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip())
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.size = Pt(11)

            # Add placeholder box
            p2 = doc.add_paragraph('[Screenshot placeholder - Insert image here]')
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.runs[0].font.italic = True
            p2.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Add border/box for screenshot
            doc.add_paragraph('_' * 80)
            continue

        # Lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            # Bullet list
            text = line.strip()[2:]
            # Handle bold and other formatting
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            continue

        # Numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            continue

        # Bold headers (lines starting with **)
        if line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(11)
            continue

        # Regular paragraphs
        if line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip())
        else:
            # Empty line
            if i > 0 and lines[i-1].strip():  # Only add space if previous line had content
                doc.add_paragraph()

    # Save the document
    output_file = 'Library_Management_System_User_Manual.docx'
    doc.save(output_file)
    print(f"✅ Word document created successfully: {output_file}")
    print(f"📄 Document is ready for editing and adding screenshots!")
    return output_file

def add_formatted_text(paragraph, text):
    """Add text to paragraph with formatting (bold, italic, code)"""
    # Handle inline code `text`
    parts = re.split(r'(`[^`]+`)', text)

    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            # Inline code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(200, 0, 0)
        else:
            # Handle bold **text**
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for bold_part in bold_parts:
                if bold_part.startswith('**') and bold_part.endswith('**'):
                    run = paragraph.add_run(bold_part[2:-2])
                    run.font.bold = True
                else:
                    # Handle checkmarks and special characters
                    text_clean = bold_part.replace('✅', '☑').replace('❌', '☒').replace('⚠️', '⚠')
                    if text_clean:
                        paragraph.add_run(text_clean)

if __name__ == '__main__':
    print("Converting USER_MANUAL.md to Word document...")
    print("-" * 60)
    create_word_document()
    print("-" * 60)
    print("You can now:")
    print("  1. Open the .docx file in Microsoft Word")
    print("  2. Add screenshots to the placeholders")
    print("  3. Adjust formatting as needed")
    print("  4. Save and print for submission")
